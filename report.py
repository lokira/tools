"""
This module includes methods for report generating.
"""
import time
from docx import Document
import os
import matplotlib.pyplot as plt
from docx.shared import RGBColor
import utilities as uti
from check_entry import EntryType
from logger import *
import mail

g_product_number = ''
g_tester = ''
g_req_file = ''
g_golden_file = ''
g_dut_file = ''
g_output_dir = ''
g_pictures = []
g_comments = []
g_commands = []

test_result = "PASSED"

result_path = ''

document = Document()
report_file = ''

def store_parameter(product_number, tester, req_file, golden_file, dut_file, output_dir):
    """
    Store the parameters to module report.
    Arguments:
        product_number - product number of DUT
        tester - tester of DB check
        req_file - requirement file
        golden_file - golden file
        dut_file - dut file
        output_dir - output directory
    """
    global g_product_number
    global g_tester
    global g_req_file
    global g_golden_file
    global g_dut_file
    global g_output_dir
    global result_path

    g_product_number = product_number
    g_tester = tester
    g_req_file = req_file
    g_golden_file = golden_file
    g_dut_file = dut_file
    g_output_dir = output_dir

    if not g_output_dir:
        g_output_dir = os.getcwd()
    result_path_suffix = time.strftime("_%Y%m%d_%H%M%S")
    result_path = os.path.join(g_output_dir, "result" + result_path_suffix)
    logger().info("result_path= %s" % result_path)
    uti.create_dir(result_path)


def save_tdata(cmdline, comments, t_data):
    """
    Save the table data array into the result array.
    Arguments:
        cmdline - command line
        comments - comments for the command
        t_data - table data array
    """
    global result_path
    global document
    global g_pictures
    global g_comments
    global g_commands

    g_pictures.append(t_data)
    g_comments.append(comments)
    g_commands.append(cmdline)
    logger().debug("Data of command(%s) is saved." % cmdline)
    if comments != "Correct":
        global test_result
        test_result = "FAILED"


def save_figure(cmdline):
    """
    Save the figure to a .png file.
    Arguments:
        cmdline - command line
    """
    global result_path

    dir_existed = os.path.isdir(result_path)
    if not dir_existed:
        uti.create_dir(result_path)
    figure_suffix = time.strftime("%Y%m%d_%H%M%S")
    filename = os.path.abspath(result_path+'\\{}_{}.png'.format(cmdline.replace('/', '_').replace(':', '-'), figure_suffix))
    plt.gcf().savefig(filename)

    logger().debug("Figure of command(%s) is saved as:%s" % (cmdline, filename))

    return filename


def add_parameter():
    """
    Add DB check parameters to the report document.
    """
    global document
    global g_product_number
    global g_tester
    global g_req_file
    global g_golden_file
    global g_dut_file

    document.add_heading('DB CHECK REPORT', level=0)
    document.add_paragraph('Product number: ' + g_product_number)
    document.add_paragraph('Tester:         ' + g_tester)

    date = time.strftime("%Y.%m.%d %H:%M")
    document.add_paragraph('DB Check Date: ' + date)
    document.add_paragraph('Requirement File:\n' + g_req_file)
    document.add_paragraph('Golden File:\n' + g_golden_file)
    document.add_paragraph('DUT File:\n' + g_dut_file)


def add_conclusion(conclusion):
    """
    Add DB check conclusion to the report document.
    """
    global document

    document.add_heading('\nConclusion', level=0)
    p = document.add_paragraph('Test Result:    ')
    if conclusion == 'PASSED':
        p.add_run(conclusion).font.color.rgb = RGBColor(0x22, 0x8B, 0x22)
        document.add_paragraph('Conclusion:     This test sw is ok to release.')
    else:
        p.add_run(conclusion).font.color.rgb = RGBColor(0xFF, 0x0, 0x0)
        document.add_paragraph('Conclusion:     This test sw is NOK to release.')


def add_table(entry_list):
    """
    Add a table of all commands with comments to the report document.
    """
    document.add_page_break()
    document.add_heading('\nThe list of commands', level=0)
    table = document.add_table(rows=len(entry_list) + 1, cols=2, style="Medium Grid 1 Accent 1")

    table.cell(0, 0).text = "Commands"
    table.cell(0, 1).text = "Comments"

    for i in range(len(entry_list)):
        entry = entry_list[i]
        table.cell(i+1, 0).text = entry.get_title()
        table.cell(i+1, 1).text = entry.get_conclusion()
        if entry.get_conclusion() == "Wrong":
            table.cell(i+1, 1).paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0x0, 0x0)
            table.cell(i+1, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0x0, 0x0)
        elif entry.is_ignored():
            table.cell(i+1, 1).paragraphs[0].runs[0].font.color.rgb = RGBColor(0x70, 0x70, 0x70)
            table.cell(i+1, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(0x70, 0x70, 0x70)


def add_pictures(entry_list):
    """
    Add all figures with commands and comments to the report document.
    """
    document.add_page_break()
    document.add_heading('Pictures', level=0)

    for entry in entry_list:
        if entry.is_ignored():   # don't save figure for ignored entry.
            continue
        document.add_paragraph('Command: ' + entry.get_title())
        document.add_paragraph('Comments: ' + entry.get_comment())
        if entry.etype == EntryType.table:
            t_data = entry.get_ref()
            table = document.add_table(rows=len(t_data), cols=3, style="Table Grid")
            for idx in range(len(t_data)):
                for j in range(3):
                    table.cell(idx, j).text = t_data[idx][j]
                if t_data[idx][2] == "NOK":
                        table.cell(idx, 2).paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0x0, 0x0)
        else:
            print(entry.get_ref())
            document.add_picture(entry.get_ref())
        document.add_page_break()
    return True


def save_report():
    """
    Save the report document.
    """
    global result_path
    global report_file
    dir_existed = os.path.isdir(result_path)
    if not dir_existed:
        uti.create_dir(result_path)

    report_file = os.path.join(result_path, "DbCheckReport.docx")
    try:
        document.save(report_file)
    except PermissionError:
        logger().exception('Report file has been opened by another program.')
        os.sys.exit(1)
    else:
        logger().info("%s is saved successfully." % os.path.abspath(report_file))


def generate_test_report(entry_list, conclusion):
    """
    Generate the DB check report document.
    """
    add_parameter()
    add_conclusion(conclusion)
    add_table(entry_list)
    add_pictures(entry_list)
    save_report()
    mail.send_mail(entry_list, report_file)

