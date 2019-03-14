import time
from docx import Document
import os
import matplotlib.pyplot as plt
import MyPicture as mp
from docx.shared import RGBColor

result_path = "output"
index = time.strftime("%Y%m%d_%H%M%S")
date = time.strftime("%Y.%m.%d %H:%M")
document = Document()

g_pictures = []
g_comments = []
g_commands = []


def create_report(product_number, tester, req_file, golden_file, dut_file):
    global result_path
    global index
    global document

    dir_existed = os.path.isdir(result_path)
    if not dir_existed:
        os.mkdir(result_path)

    document.add_heading('DB CHECK REPORT', level=0)

    document.add_paragraph('Product number: ' + product_number)
    document.add_paragraph('Tester:         ' + tester)
    global date
    document.add_paragraph('DB Check Date: ' + date)
    document.add_paragraph('Requirement File:\n' + req_file)
    document.add_paragraph('Golden File:\n' + golden_file)
    document.add_paragraph('DUT File:\n' + dut_file)

    try:
        document.save(os.path.abspath(result_path + '\\DbCheckReport.docx'))
    except PermissionError:
        print('Report file has been opened by another program.')
        os.sys.exit(1)
    else:
        print("Report is created successfully.")
    return True


def save_output(cmdline):
    global result_path
    global index
    global document

    dir_existed = os.path.isdir(result_path)
    if not dir_existed:
        os.mkdir(result_path)
    filename = os.path.abspath(result_path+'\\{}_{}.png'.format(cmdline.replace('/', '_'), index))

    plt.savefig(filename)
    mp.show_picture(filename)

    print(filename)

    global g_pictures
    global g_comments
    global g_commands
    g_pictures.append(filename)
    g_comments.append(mp.comments)
    g_commands.append(cmdline)
    #document.add_picture(filename)
    #document.add_paragraph('Comments: ' + mp.comments)
    #document.save(os.path.abspath(result_path + '\\DbCheckReport.docx'))


def add_conclusion():
    global result_path
    global index
    global document

    dir_existed = os.path.isdir(result_path)
    if not dir_existed:
        os.mkdir(result_path)

    document.add_heading('\nConclusion', level=0)
    p = document.add_paragraph('Test Result:    ')
    if mp.test_result == 1:
        p.add_run('PASSED').font.color.rgb = RGBColor(0x22, 0x8B, 0x22)
        document.add_paragraph('Conclusion:     This test sw is ok to release.')
    else:
        p.add_run('FAILED').font.color.rgb = RGBColor(0xFF, 0x0, 0x0)
        document.add_paragraph('Conclusion:     This test sw is NOK to release.')
    document.save(os.path.abspath(result_path + '\\DbCheckReport.docx'))


def add_table():
    global g_commands
    global g_comments

    if len(g_commands) != len(g_comments):
        print('There is error when adding table in report.')
        return False

    document.add_heading('\nResult for the commands', level=0)
    row_num = len(g_commands)
    table = document.add_table(rows=row_num + 1, cols=2, style="Medium Grid 1 Accent 1")

    table.cell(0, 0).text = "Commands"
    table.cell(0, 1).text = "Comments"

    for row in range(1, row_num + 1):
        table.cell(row, 0).text = g_commands[row - 1]
        table.cell(row, 1).text = g_comments[row - 1]


def add_pictures():
    global g_pictures
    global g_comments
    global result_path
    global g_commands

    document.add_page_break()
    document.add_heading('Pictures', level=0)

    if len(g_pictures) != len(g_comments):
        print('There is error when drawing picture.')
        return False

    for i in range(0, len(g_pictures)):
        document.add_paragraph('Command: ' + g_commands[i])
        document.add_paragraph('Comments: ' + g_comments[i])
        document.add_picture(g_pictures[i])
        document.add_page_break()
        document.save(os.path.abspath(result_path + '\\DbCheckReport.docx'))
    return True


